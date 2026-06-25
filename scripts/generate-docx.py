"""
Génère les fiches pratiques en format .docx
Usage: python3 scripts/generate-docx.py
Output: public/resources/*.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), '..', 'public', 'resources')
os.makedirs(OUT, exist_ok=True)

CORAL = RGBColor(0xE6, 0x39, 0x48)
NAVY  = RGBColor(0x2E, 0x23, 0x48)
SLATE = RGBColor(0x5E, 0x6B, 0x7D)
GREY  = RGBColor(0xA0, 0x98, 0x90)
WARN  = RGBColor(0x8B, 0x50, 0x10)


def new_doc(title: str, subtitle: str) -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # En-tête : marchépublic.be uniquement
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hdr.add_run("marchépublic.be")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = CORAL

    r2 = hdr.add_run("  ·  Ressource pratique — document de travail à adapter")
    r2.font.size = Pt(9)
    r2.font.color.rgb = SLATE

    # Séparateur coral
    sep = doc.add_paragraph()
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E63948')
    pBdr.append(bottom)
    pPr.append(pBdr)
    sep.paragraph_format.space_after = Pt(4)

    # Avertissement de conformité en haut de chaque document
    warn_box = doc.add_paragraph()
    warn_box.paragraph_format.space_before = Pt(6)
    warn_box.paragraph_format.space_after  = Pt(10)
    rw = warn_box.add_run(
        "⚠  DOCUMENT DE TRAVAIL — À ADAPTER À VOTRE CONTEXTE\n"
        "Ce document est une aide pratique éducative. Il ne constitue pas un avis juridique "
        "et ne garantit pas la conformité légale de votre procédure. "
        "Avant tout usage dans une procédure réelle, faites-le valider par une personne compétente "
        "en marchés publics (juriste, conseiller juridique, bailleur de fonds). "
        "marchépublic.be est un outil indépendant — non affilié à une autorité publique belge."
    )
    rw.font.size = Pt(8.5)
    rw.font.color.rgb = WARN
    rw.italic = True

    # Titre principal
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY

    # Sous-titre
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(11)
    r2.font.color.rgb = SLATE

    return doc


def heading(doc: Document, text: str, level: int = 2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13) if level == 2 else Pt(11)
    r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), 'E63948')
    pBdr.append(left)
    pPr.append(pBdr)


def body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x30, 0x28, 0x40)


def bullet(doc: Document, items: list, checked: bool = False):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        prefix = "☐  " if checked else "•  "
        r = p.add_run(prefix + item)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x30, 0x28, 0x40)


def table_grid(doc: Document, headers: list, rows: list):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E2348')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        row_cells = t.rows[ri + 1].cells
        for ci, cell in enumerate(row):
            row_cells[ci].text = cell
            run = row_cells[ci].paragraphs[0].runs[0] if row_cells[ci].paragraphs[0].runs else row_cells[ci].paragraphs[0].add_run(cell)
            run.font.size = Pt(9)
    doc.add_paragraph()


def disclaimer(doc: Document):
    """Pied de page avec avertissement légal renforcé et sources officielles."""
    doc.add_paragraph()
    sep = doc.add_paragraph()
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'E4D9CC')
    pBdr.append(top)
    pPr.append(pBdr)

    r = sep.add_run(
        "⚠  Avertissement — Ce document est fourni à titre éducatif et pratique uniquement. "
        "Il ne constitue pas un avis juridique. Il ne garantit pas la conformité de votre procédure "
        "avec la législation applicable. Adaptez-le impérativement à votre contexte avant tout usage réel. "
        "À faire valider par une personne compétente avant usage dans une procédure réelle. "
        "marchépublic.be est un outil indépendant — non affilié à une autorité publique belge.\n"
        "Sources officielles à consulter : BOSA (bosa.belgium.be) · "
        "e-Procurement (publicprocurement.be) · SPF Économie (economie.fgov.be) · "
        "Marchés publics Wallonie (marchespublics.wallonie.be) · "
        "Règlement général UE 2016/679 (RGPD) pour les données personnelles."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = SLATE
    r.italic = True


# ─── Document 1 ─────────────────────────────────────────────────────────────

def doc_sous_30k():
    doc = new_doc(
        "Acheter sous 30 000 € HTVA",
        "Les bons réflexes pour commander en règle quand le montant reste sous le seuil formel."
    )

    heading(doc, "Pourquoi ce seuil est important")
    body(doc, "En Belgique, les marchés de faible montant (estimés sous 30 000 € HTVA pour les fournitures et services) bénéficient d'un régime allégé. Aucune publication n'est obligatoire, et la procédure est simplifiée — mais cela ne signifie pas qu'aucune règle ne s'applique.")
    body(doc, "Si votre organisation est un pouvoir adjudicateur, les principes de transparence, d'égalité de traitement et de concurrence loyale restent d'application, même en dessous du seuil.")

    heading(doc, "La méthode des 3 devis")
    body(doc, "Pour les achats sous 30 000 € HTVA, la bonne pratique recommandée est de solliciter au minimum 3 offres comparatives. Voici comment procéder :")
    bullet(doc, [
        "Rédigez une description claire de votre besoin (ce que vous achetez, dans quel délai, selon quels critères).",
        "Envoyez la même demande à au moins 3 prestataires différents.",
        "Comparez les offres sur des critères définis à l'avance (prix, délai, qualité, références).",
        "Documentez votre choix et conservez les offres reçues.",
        "Signez le contrat ou la commande après avoir justifié votre sélection par écrit.",
    ])

    heading(doc, "Tableau comparatif simplifié")
    table_grid(doc,
        ["Critère", "Prestataire A", "Prestataire B", "Prestataire C"],
        [
            ["Prix HTVA", "", "", ""],
            ["Délai de livraison", "", "", ""],
            ["Références similaires", "", "", ""],
            ["Qualité de l'offre", "", "", ""],
            ["Score global /10", "", "", ""],
        ]
    )

    heading(doc, "Ce qu'il faut conserver dans votre dossier")
    bullet(doc, [
        "La description du besoin envoyée aux prestataires.",
        "Les 3 offres reçues (mails, devis, propositions).",
        "Le tableau comparatif complété.",
        "La justification écrite de votre choix final.",
        "Le bon de commande ou contrat signé.",
        "Les éventuelles factures et preuves de paiement.",
    ], checked=True)

    heading(doc, "Attention : quand la règle change")
    body(doc, "Ce régime allégé ne s'applique pas si :")
    bullet(doc, [
        "Votre convention de subvention impose explicitement le respect des marchés publics.",
        "Le montant total de la commande (ou des commandes similaires sur l'année) dépasse 30 000 € HTVA.",
        "Vous fractionnez artificiellement un marché pour rester sous le seuil (c'est interdit).",
    ])

    disclaimer(doc)
    doc.save(os.path.join(OUT, 'acheter-sous-30000.docx'))
    print("✓ acheter-sous-30000.docx")


# ─── Document 2 ─────────────────────────────────────────────────────────────

def doc_comparer():
    doc = new_doc(
        "Comparer plusieurs prestataires simplement",
        "Une méthode claire pour mettre en concurrence sans alourdir votre projet."
    )

    heading(doc, "Pourquoi formaliser la comparaison ?")
    body(doc, "Choisir un prestataire de façon équitable et traçable protège votre organisation. En cas de contrôle par votre bailleur de fonds, un audit ou un recours d'un prestataire écarté, vous devez pouvoir justifier votre décision.")

    heading(doc, "Méthode en 4 étapes")
    bullet(doc, [
        "Étape 1 — Définir le besoin : rédigez une fiche de 1 à 2 pages décrivant ce que vous attendez, les livrables, le budget indicatif et les délais.",
        "Étape 2 — Identifier les prestataires : cherchez au moins 3 candidats via votre réseau, des plateformes sectorielles ou des recommandations.",
        "Étape 3 — Envoyer la même demande : transmettez exactement le même brief à chaque prestataire pour permettre une comparaison équitable.",
        "Étape 4 — Évaluer et décider : utilisez la grille ci-dessous pour noter chaque offre et justifier votre choix final.",
    ])

    heading(doc, "Grille d'évaluation pondérée")
    table_grid(doc,
        ["Critère", "Poids (%)", "Prestataire A /10", "Prestataire B /10", "Prestataire C /10"],
        [
            ["Compréhension du besoin", "25 %", "", "", ""],
            ["Prix total HTVA", "30 %", "", "", ""],
            ["Références et expérience", "20 %", "", "", ""],
            ["Délai et disponibilité", "15 %", "", "", ""],
            ["Qualité de la proposition", "10 %", "", "", ""],
            ["TOTAL PONDÉRÉ /10", "100 %", "", "", ""],
        ]
    )

    heading(doc, "Modèle de justification du choix")
    body(doc, "À compléter et à conserver dans votre dossier :")
    bullet(doc, [
        "Prestataire retenu : ___________________________",
        "Score obtenu : ________ / 10",
        "Raison principale du choix : ___________________________",
        "Date de la décision : ___________________________",
        "Signataire : ___________________________",
    ])

    heading(doc, "Checklist finale avant signature")
    bullet(doc, [
        "Les 3 offres sont conservées dans votre dossier.",
        "La grille d'évaluation est complétée et signée.",
        "Le choix est justifié par écrit.",
        "Le contrat ou bon de commande est prêt.",
        "Votre bailleur de fonds a été informé si requis.",
    ], checked=True)

    disclaimer(doc)
    doc.save(os.path.join(OUT, 'comparer-prestataires.docx'))
    print("✓ comparer-prestataires.docx")


# ─── Document 3 ─────────────────────────────────────────────────────────────

def doc_cadrer_digital():
    doc = new_doc(
        "Cadrer un projet digital avant de demander des prix",
        "Définir votre besoin web ou logiciel pour obtenir des offres comparables et pertinentes."
    )

    heading(doc, "Pourquoi cadrer avant de prospecter ?")
    body(doc, "Si vous contactez des agences ou développeurs sans avoir clarifié votre besoin, vous recevrez des offres incomparables — certaines sous-estimées, d'autres surestimées. Un brief clair permet d'obtenir des prix justes et des propositions adaptées.")

    heading(doc, "Les 6 éléments d'un bon brief digital")
    bullet(doc, [
        "1. Contexte : qui êtes-vous, quel est votre public cible, pourquoi ce projet maintenant ?",
        "2. Objectifs : qu'est-ce que ce projet doit accomplir concrètement ? (ex. : augmenter les inscriptions, faciliter la gestion interne)",
        "3. Périmètre fonctionnel : liste des fonctionnalités attendues (formulaire, espace membres, paiement en ligne…)",
        "4. Contraintes techniques : hébergement existant, CMS préféré, intégrations requises (CRM, newsletter…)",
        "5. Budget indicatif : une fourchette réaliste permet d'obtenir des offres adaptées.",
        "6. Délais : date de livraison souhaitée, jalons intermédiaires si nécessaire.",
    ])

    heading(doc, "Exemple de structure de brief")
    table_grid(doc,
        ["Élément", "Vos réponses"],
        [
            ["Nom de votre organisation", ""],
            ["Type de projet (site vitrine, appli, outil…)", ""],
            ["Public cible", ""],
            ["Fonctionnalités principales (liste)", ""],
            ["Fonctionnalités secondaires (liste)", ""],
            ["CMS souhaité (WordPress, autre…)", ""],
            ["Intégrations requises", ""],
            ["Budget indicatif HTVA", ""],
            ["Date de livraison souhaitée", ""],
            ["Interlocuteur principal", ""],
        ]
    )

    heading(doc, "Ce que vous devez demander à chaque prestataire")
    bullet(doc, [
        "Une proposition détaillée avec ventilation des postes de coût.",
        "Les technologies utilisées et pourquoi.",
        "Le mode de livraison (jalons, recettage, formation).",
        "La maintenance post-livraison (incluse ou en option ?).",
        "Des références de projets similaires réalisés.",
        "Les conditions de propriété du code source ou du site.",
    ])

    heading(doc, "Checklist avant d'envoyer votre brief")
    bullet(doc, [
        "Le brief tient en 1 à 3 pages maximum.",
        "Les fonctionnalités sont listées par priorité (must-have / nice-to-have).",
        "Le budget indicatif est mentionné.",
        "La date de réponse souhaitée est précisée.",
        "Le brief est identique pour tous les prestataires contactés.",
    ], checked=True)

    disclaimer(doc)
    doc.save(os.path.join(OUT, 'cadrer-projet-digital.docx'))
    print("✓ cadrer-projet-digital.docx")


# ─── Document 4 ─────────────────────────────────────────────────────────────

def doc_over_30k():
    doc = new_doc(
        "Au-delà de 30 000 € HTVA — Que faire ?",
        "Les procédures qui se déclenchent au-delà du seuil et comment les aborder sereinement."
    )

    heading(doc, "Tableau des seuils (fournitures et services)")
    table_grid(doc,
        ["Montant estimé HTVA", "Régime applicable", "Publication requise"],
        [
            ["< 30 000 €", "Régime allégé — 3 devis recommandés", "Non"],
            ["30 000 € – 143 000 €", "Procédure négociée sans publication", "Non (facultatif)"],
            ["143 000 € – 221 000 €", "Procédure négociée avec publication", "e-Procurement"],
            ["> 221 000 €", "Procédure ouverte (appel d'offres)", "e-Procurement + JOUE"],
        ]
    )
    body(doc, "Note : pour les marchés de travaux, les seuils sont différents (seuil européen à 5 382 000 €). Ces montants peuvent être révisés. Vérifiez les seuils actuels sur bosa.belgium.be avant toute procédure.")

    heading(doc, "Procédure négociée sans publication (30k–143k €)")
    body(doc, "C'est la procédure la plus fréquente pour les ASBL soumises aux règles. Elle comprend :")
    bullet(doc, [
        "Rédaction d'un cahier des charges simplifié (description du besoin, critères d'attribution).",
        "Invitation d'au minimum 3 prestataires à remettre une offre.",
        "Délai de remise des offres : minimum 10 jours calendrier (recommandé).",
        "Évaluation des offres sur les critères définis.",
        "Information des candidats non retenus.",
        "Signature du marché et conservation du dossier.",
    ])

    heading(doc, "Documents à préparer")
    bullet(doc, [
        "Cahier spécial des charges (CSC) ou document descriptif.",
        "Formulaire d'invitation à soumissionner.",
        "Grille d'évaluation des offres.",
        "Procès-verbal d'ouverture et d'analyse des offres.",
        "Lettre de notification au prestataire retenu.",
        "Lettre d'information aux prestataires écartés.",
        "Contrat ou bon de commande signé.",
    ], checked=True)

    heading(doc, "Ressources officielles à consulter")
    bullet(doc, [
        "BOSA : bosa.belgium.be/fr/marche-publics-regles-et-procedures",
        "SPF Économie : economie.fgov.be/fr/themes/marches-publics",
        "Portail Wallonie : marchespublics.wallonie.be",
        "e-Procurement : publicprocurement.be",
    ])

    disclaimer(doc)
    doc.save(os.path.join(OUT, 'au-dela-de-30000.docx'))
    print("✓ au-dela-de-30000.docx")


# ─── Document 5 ─────────────────────────────────────────────────────────────

def doc_asbl_subsidiee():
    doc = new_doc(
        "ASBL subsidiée — Quelles obligations d'achat vérifier ?",
        "Les points à contrôler dans vos conventions de subvention avant de choisir un prestataire."
    )

    heading(doc, "Trois sources d'obligations possibles")
    bullet(doc, [
        "1. La loi du 17 juin 2016 sur les marchés publics — si vous êtes qualifié·e de pouvoir adjudicateur.",
        "2. Votre convention de subvention — certains bailleurs imposent contractuellement des règles de mise en concurrence.",
        "3. Les règlements internes de votre bailleur — certains subsidiant ont leurs propres procédures.",
    ])
    body(doc, "Ces trois sources sont indépendantes : vous pouvez être soumis à l'une sans l'être aux autres. Dans le doute, vérifiez les trois.")

    heading(doc, "Ce que vous devez lire dans votre convention")
    bullet(doc, [
        "Y a-t-il une clause mentionnant 'marchés publics', 'mise en concurrence' ou 'appel d'offres' ?",
        "Y a-t-il un montant seuil au-delà duquel une procédure formelle est requise ?",
        "Le bailleur impose-t-il un nombre minimum de devis ?",
        "Y a-t-il une obligation de conserver et transmettre les pièces justificatives ?",
        "Le bailleur doit-il approuver le choix du prestataire à l'avance ?",
    ])

    heading(doc, "Checklist de vérification")
    table_grid(doc,
        ["Point à vérifier", "Oui", "Non", "À clarifier"],
        [
            ["Convention lue intégralement ?", "", "", ""],
            ["Clause marché public identifiée ?", "", "", ""],
            ["Seuil de mise en concurrence noté ?", "", "", ""],
            ["Nombre de devis requis connu ?", "", "", ""],
            ["Pièces à conserver listées ?", "", "", ""],
            ["Approbation préalable du bailleur requise ?", "", "", ""],
            ["Règlements internes du bailleur consultés ?", "", "", ""],
        ]
    )

    heading(doc, "Si vous ne trouvez pas la réponse dans la convention")
    bullet(doc, [
        "Contactez votre chargé·e de dossier chez le bailleur de fonds.",
        "Demandez une clarification par écrit (mail) et conservez la réponse.",
        "En l'absence de réponse claire, appliquez le principe de précaution : procédez à une mise en concurrence documentée.",
    ])

    heading(doc, "Ce que vous devez conserver dans votre dossier de subvention")
    bullet(doc, [
        "La convention de subvention signée.",
        "Les offres reçues (minimum 3 si exigé).",
        "Le tableau comparatif des offres.",
        "La justification du choix retenu.",
        "Le contrat avec le prestataire.",
        "Les factures et preuves de paiement.",
        "Toute correspondance avec le bailleur sur les achats.",
    ], checked=True)

    disclaimer(doc)
    doc.save(os.path.join(OUT, 'asbl-subsidiee-obligations.docx'))
    print("✓ asbl-subsidiee-obligations.docx")


# ─── Document 6 — Template demande de prix ───────────────────────────────────

def doc_template_demande_prix():
    doc = new_doc(
        "Template — Demande de prix / Consultation simplifiée",
        "Document de travail à adapter — aide pratique pour comparer des prestataires. Non validé juridiquement."
    )

    heading(doc, "Mode d'emploi de ce template")
    body(doc, "Ce document est un modèle de travail simplifié. Il vous aide à structurer une demande de prix informelle ou une consultation simplifiée selon la méthode des 3 devis. Il ne remplace pas un cahier spécial des charges (CSC) au sens de la loi du 17 juin 2016.")
    body(doc, "Adaptez chaque section à votre contexte avant d'envoyer. Si votre marché dépasse 30 000 € HTVA ou si votre convention de subvention impose une procédure formelle, utilisez ce document uniquement comme point de départ et faites-le valider.")

    heading(doc, "SECTION A — Identification de votre organisation")
    table_grid(doc,
        ["Champ", "À compléter"],
        [
            ["Nom de l'organisation", ""],
            ["Type (ASBL, fondation, autre)", ""],
            ["Adresse", ""],
            ["Contact responsable de l'achat", ""],
            ["Email de contact", ""],
            ["Date d'envoi de la demande", ""],
            ["Date limite de réponse souhaitée", ""],
        ]
    )

    heading(doc, "SECTION B — Description du besoin")
    body(doc, "Décrivez ci-dessous ce que vous souhaitez acheter ou commander. Soyez précis : le prestataire doit pouvoir faire une offre comparable aux autres candidats.")
    table_grid(doc,
        ["Élément", "Description"],
        [
            ["Objet de la demande (en 1 phrase)", ""],
            ["Description détaillée de la prestation", ""],
            ["Livrables attendus", ""],
            ["Délai de réalisation souhaité", ""],
            ["Lieu de prestation (si applicable)", ""],
            ["Budget indicatif HTVA (facultatif)", ""],
            ["Conditions particulières", ""],
        ]
    )

    heading(doc, "SECTION C — Critères d'évaluation des offres")
    body(doc, "Définissez vos critères AVANT de recevoir les offres. Ne les modifiez pas après réception.")
    table_grid(doc,
        ["Critère", "Poids (%)", "Description"],
        [
            ["Prix total HTVA", "_____ %", "Offre la plus compétitive"],
            ["Compréhension du besoin", "_____ %", "Qualité de la proposition"],
            ["Expérience / références", "_____ %", "Projets similaires réalisés"],
            ["Délai de réalisation", "_____ %", "Respect du calendrier"],
            ["Autres (à préciser)", "_____ %", ""],
            ["TOTAL", "100 %", ""],
        ]
    )

    heading(doc, "SECTION D — Tableau comparatif des offres reçues")
    body(doc, "À compléter après réception des offres. Conservez ce tableau dans votre dossier.")
    table_grid(doc,
        ["Critère", "Poids", "Prestataire A", "Prestataire B", "Prestataire C"],
        [
            ["Prix HTVA", "", "", "", ""],
            ["Score compréhension /10", "", "", "", ""],
            ["Score expérience /10", "", "", "", ""],
            ["Score délai /10", "", "", "", ""],
            ["Score autres /10", "", "", "", ""],
            ["TOTAL PONDÉRÉ /10", "", "", "", ""],
        ]
    )

    heading(doc, "SECTION E — Justification du choix")
    body(doc, "À compléter et à signer avant de notifier les prestataires.")
    bullet(doc, [
        "Prestataire retenu : ___________________________",
        "Score total obtenu : ________ / 10",
        "Raison principale du choix : ___________________________",
        "Date de la décision : ___________________________",
        "Nom et fonction du signataire : ___________________________",
        "Signature : ___________________________",
    ])

    heading(doc, "SECTION F — Checklist dossier à conserver")
    bullet(doc, [
        "Ce document complété et signé.",
        "Les 3 offres reçues (emails, devis, propositions écrites).",
        "Le tableau comparatif complété (Section D).",
        "La justification du choix (Section E).",
        "Le bon de commande ou contrat signé avec le prestataire retenu.",
        "Les éventuelles lettres d'information aux prestataires non retenus.",
        "Toute correspondance pertinente avec le bailleur de fonds.",
    ], checked=True)

    heading(doc, "Sources officielles à consulter")
    bullet(doc, [
        "BOSA — règles et procédures : bosa.belgium.be/fr/marche-publics-regles-et-procedures",
        "e-Procurement — plateforme officielle belge : publicprocurement.be",
        "SPF Économie — marchés publics : economie.fgov.be/fr/themes/marches-publics",
        "Marchés publics Wallonie : marchespublics.wallonie.be",
    ])

    disclaimer(doc)
    doc.save(os.path.join(OUT, 'template-demande-de-prix.docx'))
    print("✓ template-demande-de-prix.docx")


if __name__ == '__main__':
    doc_sous_30k()
    doc_comparer()
    doc_cadrer_digital()
    doc_over_30k()
    doc_asbl_subsidiee()
    doc_template_demande_prix()
    print("\nDone — fichiers dans public/resources/")
